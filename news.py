import requests
from bs4 import BeautifulSoup
from docx import Document

def main():
    url = input("Enter the news article URL: ").strip()
    output_file = "article_content.docx"
    try:
        response = requests.get(url)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, "html.parser")
        
        # Extract the title
        title = soup.title.string if soup.title else "No Title"
        
        # Try to extract the main headline (OnlineKhabar uses h1)
        headline = soup.find("h1")
        if headline:
            title = headline.get_text(strip=True)
        
        # Try to extract main article body (works for OnlineKhabar and many sites)
        paragraphs = []
        # OnlineKhabar uses <div class="ok-news-post-content clearfix"> for article content
        article_body = soup.find("div", class_="ok-news-post-content")
        if not article_body:
            # fallback: get all paragraphs
            paragraphs = [p.get_text(strip=True) for p in soup.find_all("p")]
        else:
            paragraphs = [p.get_text(strip=True) for p in article_body.find_all("p")]
        
        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(url)
        doc.add_paragraph('---')
        for para in paragraphs:
            doc.add_paragraph(para)
        
        doc.save(output_file)
        print(f"Article content saved to {output_file}")

    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    main()
